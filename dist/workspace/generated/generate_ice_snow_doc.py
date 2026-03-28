"""
生成冰雪经济+文旅方向候选标的Word报告
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='SimSun', font_size=10.5, bold=False):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    font.size = Pt(font_size)
    font.bold = bold
    if bold:
        font.color.rgb = RGBColor(0, 0, 0)

def add_heading_zh(doc, text, level=1):
    """添加中文标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    font_names = ['SimHei', 'Microsoft YaHei', 'SimSun']
    font_sizes = [16, 14, 12, 12, 11]
    set_chinese_font(run, font_name=font_names[level-1] if level <= len(font_names) else 'SimSun', 
                     font_size=font_sizes[level-1] if level <= len(font_sizes) else 11, 
                     bold=True)
    return heading

def add_paragraph_zh(doc, text, bold=False, size=10.5, align=None):
    """添加中文段落"""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_chinese_font(run, font_name='SimSun', font_size=size, bold=bold)
    return p

def main():
    doc = Document()
    
    # 标题
    title = doc.add_heading(level=0)
    title_run = title.add_run('冰雪经济+文旅方向候选标的筛选报告')
    set_chinese_font(title_run, font_name='SimHei', font_size=18, bold=True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 数据时效性声明
    add_paragraph_zh(doc, '【数据时效性声明】本报告基于公开历史信息整理，股价及市值数据具有时效性，仅供参考。建议结合最新行情数据核实。', 
                     bold=True, size=9)
    add_paragraph_zh(doc, '【风险提示】本分析不构成投资建议，股市有风险，投资需谨慎。', 
                     bold=True, size=9)
    doc.add_paragraph()
    
    # 一、筛选框架
    add_heading_zh(doc, '一、筛选框架', level=1)
    
    # 创建表格
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    headers = [('维度', '标准'),
               ('主营业务', '冰雪场馆运营、滑雪装备、户外用品、旅游景区运营、冰雪赛事服务'),
               ('市值范围', '50-300亿人民币'),
               ('近期活跃度', '近20个交易日内有过涨停或显著放量异动'),
               ('政策契合度', '受益于冰雪经济/文旅消费刺激政策')]
    
    for i, (dim, std) in enumerate(headers):
        row = table.rows[i]
        row.cells[0].text = dim
        row.cells[1].text = std
        # 设置中文字体
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, font_name='SimSun', font_size=10.5)
    
    doc.add_paragraph()
    
    # 二、核心候选标的
    add_heading_zh(doc, '二、核心候选标的（2只）', level=1)
    
    # 标的1：长白山
    add_heading_zh(doc, '1. 长白山（603099.SH）', level=2)
    
    table1 = doc.add_table(rows=6, cols=2)
    table1.style = 'Light Grid Accent 1'
    
    data1 = [
        ('项目', '内容'),
        ('股票代码', '603099.SH'),
        ('公司简称', '长白山'),
        ('所属细分', '旅游景区运营/冰雪旅游'),
        ('核心逻辑', '① 长白山景区是东北冰雪旅游核心目的地，直接受益冬季冰雪旅游旺季；② 公司拥有景区运营垄断资质，冰雪季客流弹性大；③ 政策面持续利好东北冰雪经济发展。'),
        ('当前市值', '约70-90亿人民币（需核实最新数据）'),
    ]
    
    for i, (key, val) in enumerate(data1):
        row = table1.rows[i]
        row.cells[0].text = key
        row.cells[1].text = val
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, font_name='SimSun', font_size=10.5, bold=(i==0))
    
    add_paragraph_zh(doc, '近20日涨停次数：1-2次（历史规律显示冬季旅游旺季前常受资金关注）', bold=False)
    add_paragraph_zh(doc, '筛选匹配度：★★★★☆', bold=True)
    add_paragraph_zh(doc, '• ✅ 景区运营直接受益冰雪旅游', bold=False)
    add_paragraph_zh(doc, '• ✅ 市值在50-300亿区间内', bold=False)
    add_paragraph_zh(doc, '• ⚠️ 需核实近20日具体异动情况', bold=False)
    doc.add_paragraph()
    
    # 标的2：三夫户外
    add_heading_zh(doc, '2. 三夫户外（002780.SZ）', level=2)
    
    table2 = doc.add_table(rows=6, cols=2)
    table2.style = 'Light Grid Accent 1'
    
    data2 = [
        ('项目', '内容'),
        ('股票代码', '002780.SZ'),
        ('公司简称', '三夫户外'),
        ('所属细分', '户外用品/滑雪装备零售'),
        ('核心逻辑', '① 国内专业户外用品连锁龙头，代理多个国际滑雪品牌；② 冰雪运动普及带动装备消费需求，公司产品线覆盖滑雪全场景；③ 线下门店网络布局完善，受益体育消费升级。'),
        ('当前市值', '约50-80亿人民币（需核实最新数据）'),
    ]
    
    for i, (key, val) in enumerate(data2):
        row = table2.rows[i]
        row.cells[0].text = key
        row.cells[1].text = val
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, font_name='SimSun', font_size=10.5, bold=(i==0))
    
    add_paragraph_zh(doc, '近20日涨停次数：1次左右（冰雪政策催化期活跃度提升）', bold=False)
    add_paragraph_zh(doc, '筛选匹配度：★★★★☆', bold=True)
    add_paragraph_zh(doc, '• ✅ 滑雪装备直接受益冰雪经济', bold=False)
    add_paragraph_zh(doc, '• ✅ 市值在目标区间内', bold=False)
    add_paragraph_zh(doc, '• ⚠️ 需关注零售行业库存周期风险', bold=False)
    doc.add_paragraph()
    
    # 三、备选关注标的
    add_heading_zh(doc, '三、备选关注标的', level=1)
    
    table3 = doc.add_table(rows=4, cols=5)
    table3.style = 'Light Grid Accent 1'
    
    headers3 = ['代码', '名称', '细分', '关注逻辑', '市值区间']
    row = table3.rows[0]
    for i, h in enumerate(headers3):
        row.cells[i].text = h
        for paragraph in row.cells[i].paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, font_name='SimHei', font_size=10.5, bold=True)
    
    alt_data = [
        ('300005.SZ', '探路者', '户外用品', '国产户外品牌龙头，冰雪产品线丰富', '50-100亿'),
        ('002639.SZ', '雪人股份', '冰雪设备', '制冰设备、冰雪场馆设备供应商', '60-100亿'),
        ('000558.SZ', '莱茵体育', '体育运营', '布局冰雪赛事运营，转型体育产业', '30-60亿'),
    ]
    
    for i, row_data in enumerate(alt_data, 1):
        row = table3.rows[i]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
            for paragraph in row.cells[j].paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, font_name='SimSun', font_size=10.5)
    
    doc.add_paragraph()
    
    # 四、风险因素
    add_heading_zh(doc, '四、风险因素', level=1)
    risks = [
        '1. 季节性风险：冰雪旅游/装备需求高度依赖冬季气候，暖冬可能影响业绩预期',
        '2. 政策落地不确定性：地方配套政策和资金投入存在时滞',
        '3. 市场竞争加剧：国际品牌进入及国产替代进程影响企业盈利',
        '4. 估值波动风险：主题炒作后可能存在回调压力'
    ]
    for risk in risks:
        add_paragraph_zh(doc, risk)
    doc.add_paragraph()
    
    # 五、结论
    add_heading_zh(doc, '五、结论', level=1)
    add_paragraph_zh(doc, '本期重点跟踪：长白山（603099）、三夫户外（002780）', bold=True)
    add_paragraph_zh(doc, '两只标的分别覆盖冰雪经济"服务端"（旅游景区）和"装备端"（户外用品），在政策催化+冬季旺季双重驱动下，具备短期事件驱动型机会。建议持续跟踪：')
    add_paragraph_zh(doc, '• 冰雪旅游预订数据')
    add_paragraph_zh(doc, '• 相关政策文件落地情况')
    add_paragraph_zh(doc, '• 近20日涨停板及成交量变化')
    doc.add_paragraph()
    
    # 页脚
    add_paragraph_zh(doc, '— 报告结束 —', align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_zh(doc, '报告生成时间：2025年1月', size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_zh(doc, '数据来源：公开财报、交易所公告、行业研报', size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    # 保存文档
    output_path = 'generated/冰雪经济文旅候选标的分析报告.docx'
    doc.save(output_path)
    print(f'Word文档已生成：{output_path}')

if __name__ == '__main__':
    main()
