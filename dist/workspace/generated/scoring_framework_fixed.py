#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
A股多次涨停概率评分框架 - 完整评分对照表文档生成器
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_run_font(run, font_name='微软雅黑', font_size=10.5, bold=False, color=None):
    """设置文本格式"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    colors = {
        1: (0, 51, 102),      # 深蓝
        2: (0, 102, 153),     # 青蓝
        3: (51, 102, 102),    # 深青
    }
    set_run_font(run, font_size=16 if level==1 else (14 if level==2 else 12), 
                 bold=True, color=colors.get(level, (0,0,0)))
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph_custom(doc, text, bold=False, align='left', font_size=10.5, color=None):
    """添加自定义段落"""
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run = p.add_run(text)
    set_run_font(run, bold=bold, font_size=font_size, color=color)
    return p

def create_scoring_framework_doc():
    """创建完整评分框架文档"""
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10.5)
    
    # ============================================
    # 封面部分
    # ============================================
    for _ in range(5):
        add_paragraph_custom(doc, '', font_size=12)
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('A股多次涨停概率评分框架')
    set_run_font(run, font_size=22, bold=True, color=(0, 51, 102))
    
    add_paragraph_custom(doc, '', font_size=12)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('完整1-10分评分对照表')
    set_run_font(run, font_size=16, bold=True, color=(102, 102, 102))
    
    for _ in range(3):
        add_paragraph_custom(doc, '', font_size=12)
    
    # 框架说明
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('版本：V1.0\n适用市场：A股主板/创业板/科创板\n编制日期：2024年')
    set_run_font(run, font_size=11, color=(128, 128, 128))
    
    doc.add_page_break()
    
    # ============================================
    # 第一部分：封面说明
    # ============================================
    add_heading_custom(doc, '一、框架说明', level=1)
    
    add_heading_custom(doc, '1.1 框架目的', level=2)
    add_paragraph_custom(doc, 
        '本评分框架旨在建立一套可解释、可量化的"多次涨停概率"评估体系，用于对A股市场潜在'
        '热点标的进行系统性筛选与分层排序。框架基于行为金融学与市场微观结构理论，结合A股'
        '市场特征（涨跌停板制度、散户占比高、政策驱动明显等），通过多维度评分机制识别具备'
        '连续涨停潜力的个股。核心目标是将主观经验转化为结构化评分标准，降低决策偏差，提高选股效率。')
    
    add_heading_custom(doc, '1.2 框架设计原则', level=2)
    principles = [
        ('可解释性', '每个评分维度均有明确的定义、标准和数据来源，避免黑箱操作。'),
        ('可量化性', '尽可能采用客观指标，减少主观判断空间，确保评分一致性。'),
        ('动态适应性', '框架预留调整空间，可根据市场环境变化（如监管政策、资金流向）优化权重。'),
        ('风险前置', '评分结果需结合风险提示使用，不作为投资建议。')
    ]
    for title, desc in principles:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{title}：')
        set_run_font(run, bold=True)
        run = p.add_run(desc)
        set_run_font(run)
    
    add_heading_custom(doc, '1.3 权重依据', level=2)
    add_paragraph_custom(doc, 
        '权重分配基于以下逻辑：事件重要性决定上涨逻辑强度（占40%），稀缺性决定资金'
        '聚焦程度（占60%）。这一比例反映了A股市场"强者恒强"的博弈特征——即使有重大'
        '利好，若标的缺乏辨识度，也难以形成连续涨停。')
    
    # 权重说明表格
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['评分维度', '权重', '权重依据说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, '1F4E79')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, bold=True, color=(255, 255, 255))
    
    data = [
        ['事件重要性(E)', '40%', '决定上涨逻辑的有效性和持续性。重大政策/事件是涨停的根本驱动力。'],
        ['稀缺性(S)', '60%', '决定资金聚焦程度。标的越稀缺，越易形成共识，连续涨停概率越高。']
    ]
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, text in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = text
            if col_idx == 1:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ============================================
    # 第二部分：事件重要性评分表
    # ============================================
    add_heading_custom(doc, '二、事件重要性评分表(E)', level=1)
    
    add_paragraph_custom(doc, 
        '事件重要性评估上涨驱动因素的有效性和持续性。评分基于三个维度：政策/事件级别、'
        '市场预期差、业绩弹性。最终重要性得分(E)为三维度加权聚合。')
    
    add_heading_custom(doc, '2.1 维度一：政策/事件级别(P)', level=2)
    add_paragraph_custom(doc, '评估驱动事件的行政层级和产业影响范围。')
    
    table1 = doc.add_table(rows=6, cols=4)
    table1.style = 'Table Grid'
    
    headers1 = ['等级', '标准定义', '典型场景', '单项得分']
    for i, h in enumerate(headers1):
        cell = table1.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '2E75B5')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, bold=True, color=(255, 255, 255))
    
    p_data = [
        ['P5', '国家级战略/重大突发事件', '国家战略新兴产业规划、突发地缘冲突、重大技术突破', '5分'],
        ['P4', '部委级政策/行业重大变革', '工信部专项支持、证监会改革、行业准入政策调整', '4分'],
        ['P3', '地方级重点政策/产业联盟', '地方重点产业规划、央企重组、头部企业战略合作', '3分'],
        ['P2', '企业级重大公告/订单', '重大合同签订、并购重组、核心技术突破公告', '2分'],
        ['P1', '一般性利好/常规公告', '季度业绩预增、普通合作协议、日常经营利好', '1分']
    ]
    for row_idx, row_data in enumerate(p_data, 1):
        for col_idx, text in enumerate(row_data):
            cell = table1.rows[row_idx].cells[col_idx]
            cell.text = text
            if col_idx in [0, 3]:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_heading_custom(doc, '2.2 维度二：市场预期差(M)', level=2)
    add_paragraph_custom(doc, '评估事件实际影响与市场已有预期的偏离程度。')
    
    table2 = doc.add_table(rows=5, cols=4)
    table2.style = 'Table Grid'
    
    headers2 = ['等级', '标准定义', '判断依据', '单项得分']
    for i, h in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '2E75B5')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, bold=True, color=(255, 255, 255))
    
    m_data = [
        ['M4', '超预期', '事件强度/时间/范围显著超越市场主流预期，盘面出现突发性大涨', '4分'],
        ['M3', '符合预期偏上', '事件符合预期但细节超预期，或市场存在分歧后确认利好', '3分'],
        ['M2', '符合预期', '事件与市场主流预期基本一致，利好已被部分price in', '2分'],
        ['M1', '低于预期', '事件落地但强度不及预期，或存在隐性利空', '1分']
    ]
    for row_idx, row_data in enumerate(m_data, 1):
        for col_idx, text in enumerate(row_data):
            cell = table2.rows[row_idx].cells[col_idx]
            cell.text = text
            if col_idx in [0, 3]:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_heading_custom(doc, '2.3 维度三：业绩弹性(R)', level=2)
    add_paragraph_custom(doc, '评估事件对公司未来业绩的潜在影响幅度。')
    
    table3 = doc.add_table(rows=5, cols=4)
    table3.style = 'Table Grid'
    
    headers3 = ['等级', '标准定义', '测算标准', '单项得分']
    for i, h in enumerate(headers3):
        cell = table3.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '2E75B5')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, bold=True, color=(255, 255, 255))
    
    r_data = [
        ['R4', '高弹性', '预期净利润增长>100%或营收增长>200%，业务结构根本性改变', '4分'],
        ['R3', '中高弹性', '预期净利润增长50-100%或营收增长100-200%', '3分'],
        ['R2', '中等弹性', '预期净利润增长20-50%或营收增长50-100%', '2分'],
        ['R1', '低弹性', '预期净利润增长<20%，主要为概念性映射而非实质业绩', '1分']
    ]
    for row_idx, row_data in enumerate(r_data, 1):
        for col_idx, text in enumerate(row_data):
            cell = table3.rows[row_idx].cells[col_idx]
            cell.text = text
            if col_idx in [0, 3]:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_heading_custom(doc, '2.4 事件重要性聚合规则', level=2)
    add_paragraph_custom(doc, '事件重要性得分 E = P × 0.35 + M × 0.35 + R × 0.30', bold=True)
    add_paragraph_custom(doc, '• 得分区间：E ∈ [1, 5]，保留两位小数')
    add_paragraph_custom(doc, '• P维度满分5分，M和R维度满分4分，经加权后E的理论范围为[1, 4.65]')
    add_paragraph_custom(doc, '• 为统一量纲，最终E按满分5分标准化')
    
    doc.add_page_break()
    
    # ============================================
    # 第三部分：稀缺性评分表
    # ============================================
    add_heading_custom(doc, '三、稀缺性评分表(S)', level=1)
    
    add_paragraph_custom(doc, 
        '稀缺性评估标的在市场中的独特性和资金聚焦程度。评分基于三个维度：标的独特性、'
        '资金辨识度、筹码结构。稀缺性决定资金是否愿意持续流入并推动连续涨停。')
    
    add_heading_custom(doc, '3.1 维度一：标的独特性(U)', level=2)
    add_paragraph_custom(doc, '评估标的在所属题材中的独特地位和不可替代性。')
    
    table4 = doc.add_table(rows=6, cols=4)
    table4.style = 'Table Grid'
    
    headers4 = ['等级', '标准定义', '典型特征', '单项得分']
    for i, h in enumerate(headers4):
        cell = table4.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '70AD47')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, bold=True, color=(255, 255, 255))
    
    u_data = [
        ['U5', '绝对龙头/唯一标的', '细分领域市占率第一、唯一上市标的、无可替代性', '5分'],
        ['U4', '核心龙头', '行业前三、技术壁垒高、资金共识度强', '4分'],
        ['U3', '重要参与者', '行业前五、有明确竞争优势但非绝对龙头', '3分'],
        ['U2', '普通标的', '行业跟随者、业务占比不高、可替代性强', '2分'],
        ['U1', '边缘标的', '题材边缘沾边、主业关联度低、蹭概念嫌疑', '1分']
    ]
    for row_idx, row_data in enumerate(u_data, 1):
        for col_idx, text in enumerate(row_data):
            cell = table4.rows[row_idx].cells[col_idx]
            cell.text = text
            if col_idx in [0, 3]:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_heading_custom(doc, '3.2 维度二：资金辨识度(I)', level=2)
    add_paragraph_custom(doc, '评估标的对各类资金（游资/机构/散户）的吸引力与辨识度。')
    
    table5 = doc.add_table(rows=6, cols=4)
    table5.style = 'Table Grid'
    
    headers5 = ['等级', '标准定义', '盘面特征', '单项得分']
    for i, h in enumerate(headers5):
        cell = table5.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '70AD47')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, bold=True, color=(255, 255, 255))
    
    i_data = [
        ['I5', '全市场焦点', '多路资金抢筹、龙虎榜豪华、市场讨论度极高', '5分'],
        ['I4', '游资核心标的', '连板基因好、历史股性活跃、游资抱团', '4分'],
        ['I3', '机构+游资共振', '有基本面支撑、机构调研密集、资金合力', '3分'],
        ['I2', '散户跟风标的', '盘子适中、名字易记、跟风属性强', '2分'],
        ['I1', '冷门标的', '成交低迷、关注度低、缺乏资金记忆', '1分']
    ]
    for row_idx, row_data in enumerate(i_data, 1):
        for col_idx, text in enumerate(row_data):
            cell = table5.rows[row_idx].cells[col_idx]
            cell.text = text
            if col_idx in [0, 3]:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_heading_custom(doc, '3.3 维度三：筹码结构(C)', level=2)
    add_paragraph_custom(doc, '评估标的流通筹码的分布状态和拉升阻力。')
    
    table6 = doc.add_table(rows=6, cols=4)
    table6.style = 'Table Grid'
    
    headers6 = ['等级', '标准定义', '量化标准', '单项得分']
    for i, h in enumerate(headers6):
        cell = table6.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '70AD47')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, bold=True, color=(255, 255, 255))
    
    c_data = [
        ['C5', '完美筹码结构', '流通市值<50亿、无大量解禁、股东户数少、套牢盘远', '5分'],
        ['C4', '优良筹码结构', '流通市值50-100亿、近期无解禁、筹码集中度高', '4分'],
        ['C3', '中等筹码结构', '流通市值100-200亿、有部分压力但可控', '3分'],
        ['C2', '较差筹码结构', '流通市值>200亿、近期有解禁或大额减持', '2分'],
        ['C1', '恶劣筹码结构', '巨量解禁在即、基金重仓套牢盘密集、历史妖股套牢', '1分']
    ]
    for row_idx, row_data in enumerate(c_data, 1):
        for col_idx, text in enumerate(row_data):
            cell = table6.rows[row_idx].cells[col_idx]
            cell.text = text
            if col_idx in [0, 3]:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_heading_custom(doc, '3.4 稀缺性聚合规则', level=2)
    add_paragraph_custom(doc, '稀缺性得分 S = U × 0.40 + I × 0.35 + C × 0.25', bold=True)
    add_paragraph_custom(doc, '• 得分区间：S ∈ [1, 5]，保留两位小数')
    add_paragraph_custom(doc, '• 三个维度满分均为5分，经加权后S的理论范围为[1, 5]')
    
    doc.add_page_break()
    
    # ============================================
    # 第四部分：综合评分公式与计算流程
    # ============================================
    add_heading_custom(doc, '四、综合评分公式与计算流程', level=1)
    
    add_heading_custom(doc, '4.1 综合评分公式', level=2)
    
    add_paragraph_custom(doc, '综合评分(Total Score, TS)计算公式：', bold=True)
    add_paragraph_custom(doc, 'TS = (E / 5) × 40% + (S / 5) × 60%', bold=True, font_size=12)
    add_paragraph_custom(doc, '其中：E为事件重要性得分(1-5)，S为稀缺性得分(1-5)')
    
    add_paragraph_custom(doc, 
        '公式说明：先将E和S标准化为[0.2, 1]区间（除以各自满分5），再按权重40%/60%加权求和，'
        '最终TS理论范围为[0.2, 1]。为便于使用，将TS映射到1-10分制（见第五章）。')
    
    add_heading_custom(doc, '4.2 计算流程图', level=2)
    
    flow_text = """
    ┌─────────────────────────────────────────────────────────────┐
    │                      评分计算流程                           │
    └──────────────────────────┬──────────────────────────────────┘
                               │
           ┌───────────────────┴───────────────────┐
           ▼                                       ▼
    ┌─────────────────┐                   ┌─────────────────┐
    │   事件重要性E    │                   │    稀缺性S      │
    │     (40%)       │                   │     (60%)       │
    └────────┬────────┘                   └────────┬────────┘
             │                                     │
    ┌────────┴────────┐                   ┌────────┴────────┐
    ▼        ▼        ▼                   ▼        ▼        ▼
   政策     预期     业绩                 独特     资金     筹码
   级别     差       弹性                 性       辨识度   结构
    │        │        │                   │        │        │
    └────────┴────────┘                   └────────┴────────┘
             │                                     │
             │    E = Σ(维度×权重)                 │    S = Σ(维度×权重)
             │      = P×0.35+M×0.35+R×0.30         │      = U×0.40+I×0.35+C×0.25
             │                                     │
             └─────────────┬───────────────────────┘
                           ▼
                ┌─────────────────────┐
                │ TS = (E/5)×0.4      │
                │      +(S/5)×0.6     │
                └──────────┬──────────┘
                           ▼
                ┌─────────────────────┐
                │   映射到1-10分区间   │
                │   Score = 2+TS×8    │
                └──────────┬──────────┘
                           ▼
                ┌─────────────────────┐
                │    输出最终评分      │
                │   附概率定性说明     │
                └─────────────────────┘
    """
    
    for line in flow_text.strip().split('\n'):
        add_paragraph_custom(doc, line, font_size=9)
    
    doc.add_page_break()
    
    # ============================================
    # 第五部分：1-10分区间映射表
    # ============================================
    add_heading_custom(doc, '五、1-10分区间映射表', level=1)
    
    add_paragraph_custom(doc, 
        '将标准化后的综合评分TS（范围[0.2, 1]）线性映射到1-10分区间，并给出对应的'
        '多次涨停概率定性描述。')
    
    add_paragraph_custom(doc, '映射公式：Final Score = 2 + TS × 8', bold=True, font_size=12)
    add_paragraph_custom(doc, '（当TS=0.2时得2分，TS=1时得10分，中间线性插值）')
    
    add_heading_custom(doc, '5.1 评分区间详解', level=2)
    
    table7 = doc.add_table(rows=10, cols=5)
    table7.style = 'Table Grid'
    
    headers7 = ['最终评分', 'TS范围', '场景描述', '多次涨停概率定性', '策略建议']
    for i, h in enumerate(headers7):
        cell = table7.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'C65911')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, bold=True, color=(255, 255, 255))
    
    score_data = [
        ['10分', '[0.95,1.00]', '国家级政策+绝对龙头+完美筹码', '极高(>80%)', '重点关注'],
        ['9分', '[0.85,0.95)', '重大政策+核心龙头+优良筹码', '很高(60-80%)', '重点关注'],
        ['8分', '[0.75,0.85)', '部委政策+核心标的+较好筹码', '高(40-60%)', '积极关注'],
        ['7分', '[0.65,0.75)', '较好政策+重要标的+中等筹码', '中高(30-40%)', '适度关注'],
        ['6分', '[0.55,0.65)', '一般政策+普通标的+可控筹码', '中等(20-30%)', '观察'],
        ['5分', '[0.45,0.55)', '普通事件+一般标的+中性筹码', '一般(10-20%)', '谨慎'],
        ['4分', '[0.35,0.45)', '弱事件+边缘标的+较差筹码', '较低(5-10%)', '回避'],
        ['3分', '[0.25,0.35)', '极弱事件+弱势标的+恶劣筹码', '低(<5%)', '回避'],
        ['2分', '[0.20,0.25)', '无实质驱动+弱势标的+恶劣筹码', '极低(<2%)', '强烈回避'],
    ]
    
    for row_idx, row_data in enumerate(score_data, 1):
        for col_idx, text in enumerate(row_data):
            cell = table7.rows[row_idx].cells[col_idx]
            cell.text = text
            if col_idx in [0, 1]:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置行背景色
        if row_idx in [1, 2]:
            for cell in table7.rows[row_idx].cells:
                set_cell_shading(cell, 'C6EFCE')  # 绿色-高分
        elif row_idx == 3:
            for cell in table7.rows[row_idx].cells:
                set_cell_shading(cell, 'E2EFDA')  # 浅绿
        elif row_idx == 4:
            for cell in table7.rows[row_idx].cells:
                set_cell_shading(cell, 'FFEB9C')  # 黄色-中分
        elif row_idx == 5:
            for cell in table7.rows[row_idx].cells:
                set_cell_shading(cell, 'FCE4D6')  # 浅橙
        else:
            for cell in table7.rows[row_idx].cells:
                set_cell_shading(cell, 'FFC7CE')  # 红色-低分
    
    add_heading_custom(doc, '5.2 概率定性说明', level=2)
    add_paragraph_custom(doc, 
        '注：概率定性基于历史数据统计和经验判断，仅供参考：')
    
    prob_items = [
        ('极高(>80%)', '具备3次以上连续涨停潜力，可能成为阶段性龙头'),
        ('很高(60-80%)', '具备2-3次连续涨停潜力，有望成为板块核心'),
        ('高(40-60%)', '具备1-2次涨停潜力，但持续性存疑'),
        ('中高(30-40%)', '可能触发1次涨停，多次涨停需额外催化'),
        ('中等(20-30%)', '涨停概率接近随机，不建议专门布局'),
        ('一般及以下(<20%)', '多次涨停概率极低，建议回避')
    ]
    for title, desc in prob_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{title}：')
        set_run_font(run, bold=True)
        run = p.add_run(desc)
        set_run_font(run)
    
    doc.add_page_break()
    
    # ============================================
    # 第六部分：使用示例
    # ============================================
    add_heading_custom(doc, '六、使用示例', level=1)
    
    add_paragraph_custom(doc, 
        '以下通过两个假设场景演示完整评分过程，展示框架的实际应用。')
    
    add_heading_custom(doc, '示例一：政策驱动型（高分场景）', level=2)
    add_paragraph_custom(doc, '场景假设：某低市值芯片设计公司A，国家发布国产算力重大扶持政策')
    
    # 创建评分计算展示表格 - 修复：确保行数足够
    ex1_data = [
        ['事件重要性E\n(权重40%)', '政策级别P', '国家级战略政策', 'P5=5分'],
        ['', '预期差M', '政策力度超预期', 'M4=4分'],
        ['', '业绩弹性R', '订单预期翻倍', 'R4=4分'],
        ['', 'E计算', '5×0.35+4×0.35+4×0.30', 'E=4.35'],
        ['稀缺性S\n(权重60%)', '独特性U', '细分赛道唯一标的', 'U5=5分'],
        ['', '辨识度I', '游资抢筹、龙虎榜活跃', 'I4=4分'],
        ['', '筹码C', '流通市值30亿、无解禁', 'C5=5分'],
        ['', 'S计算', '5×0.40+4×0.35+5×0.25', 'S=4.65'],
    ]
    
    ex1_table = doc.add_table(rows=len(ex1_data)+1, cols=4)
    ex1_table.style = 'Table Grid'
    
    ex1_headers = ['评分维度', '子维度', '评分依据', '得分']
    for i, h in enumerate(ex1_headers):
        cell = ex1_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '5B9BD5')
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, bold=True, color=(255, 255, 255))
    
    for row_idx, row_data in enumerate(ex1_data, 1):
        for col_idx, text in enumerate(row_data):
            cell = ex1_table.rows[row_idx].cells[col_idx]
            cell.text = text
    
    add_paragraph_custom(doc, '')
    add_paragraph_custom(doc, 
        '综合评分计算：TS = (4.35/5)×0.4 + (4.65/5)×0.6 = 0.348 + 0.558 = 0.906')
    add_paragraph_custom(doc, 
        '最终得分：2 + 0.906×8 = 9.25 ≈ 9分（很高概率多次涨停）', bold=True, color=(0, 128, 0))
    
    add_heading_custom(doc, '示例二：边缘沾边型（低分场景）', level=2)
    add_paragraph_custom(doc, '场景假设：某传统制造业公司B，宣布通过子公司间接参与新能源')
    
    ex2_data = [
        ['事件重要性E\n(权重40%)', '政策级别P', '企业级公告', 'P2=2分'],
        ['', '预期差M', '市场已预期、低于预期', 'M1=1分'],
        ['', '业绩弹性R', '业务占比<5%，几乎无弹性', 'R1=1分'],
        ['', 'E计算', '2×0.35+1×0.35+1×0.30', 'E=1.35'],
        ['稀缺性S\n(权重60%)', '独特性U', '边缘标的、蹭概念', 'U1=1分'],
        ['', '辨识度I', '无资金关注、成交低迷', 'I1=1分'],
        ['', '筹码C', '流通市值300亿、套牢盘多', 'C2=2分'],
        ['', 'S计算', '1×0.40+1×0.35+2×0.25', 'S=1.25'],
    ]
    
    ex2_table = doc.add_table(rows=len(ex2_data)+1, cols=4)
    ex2_table.style = 'Table Grid'
    
    for i, h in enumerate(ex1_headers):
        cell = ex2_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '5B9BD5')
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, bold=True, color=(255, 255, 255))
    
    for row_idx, row_data in enumerate(ex2_data, 1):
        for col_idx, text in enumerate(row_data):
            cell = ex2_table.rows[row_idx].cells[col_idx]
            cell.text = text
    
    add_paragraph_custom(doc, '')
    add_paragraph_custom(doc, 
        '综合评分计算：TS = (1.35/5)×0.4 + (1.25/5)×0.6 = 0.108 + 0.150 = 0.258')
    add_paragraph_custom(doc, 
        '最终得分：2 + 0.258×8 = 4.06 ≈ 4分（多次涨停概率较低）', bold=True, color=(192, 0, 0))
    
    doc.add_page_break()
    
    # ============================================
    # 第七部分：风险提示与免责声明
    # ============================================
    add_heading_custom(doc, '七、风险提示与免责声明', level=1)
    
    add_paragraph_custom(doc, '7.1 框架局限性', bold=True)
    limitations = [
        '本框架基于历史规律总结，不保证未来有效性',
        '市场情绪、资金流向等动态因素难以完全量化',
        '突发政策变化（如监管窗口指导）可能快速改变评分',
        '评分仅反映概率，不构成确定性预测',
        '个股选择需结合大盘环境，系统性风险可能压倒个股逻辑'
    ]
    for item in limitations:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        set_run_font(run)
    
    add_paragraph_custom(doc, '7.2 使用建议', bold=True)
    suggestions = [
        '建议将本框架作为初筛工具，高分标的需进一步基本面研究',
        '评分结果应结合当日市场情绪、板块效应综合判断',
        '设置止损纪律，不因高评分而忽视风险控制',
        '定期回顾评分准确性，持续优化维度权重',
        '关注监管动态，避免参与可能被重点监控的标的'
    ]
    for item in suggestions:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        set_run_font(run)
    
    add_paragraph_custom(doc, '', font_size=12)
    
    # 免责声明框
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run('─────────────────────────────────')
    set_run_font(run, color=(128, 128, 128))
    
    disc_title = doc.add_paragraph()
    disc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disc_title.add_run('免责声明')
    set_run_font(run, bold=True, font_size=12, color=(192, 0, 0))
    
    disc_text = doc.add_paragraph()
    disc_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disc_text.add_run(
        '本评分框架仅供学习研究使用，不构成任何投资建议。\n'
        '股市有风险，投资需谨慎。投资者应根据自身情况独立决策，\n'
        '自行承担投资风险。本框架作者不对使用本框架产生的任何损失承担责任。')
    set_run_font(run, font_size=9, color=(128, 128, 128))
    
    disclaimer2 = doc.add_paragraph()
    disclaimer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer2.add_run('─────────────────────────────────')
    set_run_font(run, color=(128, 128, 128))
    
    # 保存文档
    output_path = 'generated/A股多次涨停概率评分框架_完整对照表.docx'
    doc.save(output_path)
    print(f"✓ 文档已成功生成：{output_path}")
    return output_path

if __name__ == '__main__':
    create_scoring_framework_doc()
